#!/usr/bin/env python3
"""
OCR TIF Files to Word Document using EasyOCR
Scans all TIF files in a directory and creates a comprehensive Word document
"""

import os
import sys
from pathlib import Path
from datetime import datetime

def install_packages():
    """Install required packages"""
    print("Installing required packages...")
    os.system("pip install python-docx pillow easyocr -q")

def ocr_tif_file(reader, file_path):
    """Extract text from a TIF file using EasyOCR"""
    try:
        results = reader.readtext(file_path)
        text = '\n'.join([result[1] for result in results])
        return text.strip()
    except Exception as e:
        return f"[Error processing file: {str(e)}]"

def create_word_document(tif_folder, output_path):
    """Create a Word document from OCR'd TIF files"""
    
    try:
        from PIL import Image
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import easyocr
    except ImportError:
        install_packages()
        from PIL import Image
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import easyocr
    
    # Get all TIF files
    tif_files = sorted([f for f in os.listdir(tif_folder) if f.lower().endswith('.tif')])
    
    if not tif_files:
        print("No TIF files found!")
        return False
    
    print(f"Found {len(tif_files)} TIF files to process")
    
    # Initialize OCR reader
    print("Initializing EasyOCR reader (Arabic + English)...")
    reader = easyocr.Reader(['ar', 'en'], gpu=False)
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('OCR Document Scan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    meta_section = doc.add_paragraph()
    meta_section.add_run("Source Folder: ").bold = True
    meta_section.add_run(str(tif_folder))
    
    meta_section = doc.add_paragraph()
    meta_section.add_run("Total Pages: ").bold = True
    meta_section.add_run(str(len(tif_files)))
    
    meta_section = doc.add_paragraph()
    meta_section.add_run("OCR Method: ").bold = True
    meta_section.add_run("EasyOCR (Arabic + English)")
    
    meta_section = doc.add_paragraph()
    meta_section.add_run("Generated: ").bold = True
    meta_section.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    
    # Process each TIF file
    for idx, tif_file in enumerate(tif_files, 1):
        file_path = os.path.join(tif_folder, tif_file)
        print(f"[{idx}/{len(tif_files)}] Processing: {tif_file}...", end=" ", flush=True)
        
        # Add page heading
        page_heading = doc.add_heading(f'Page {idx}: {tif_file}', level=1)
        page_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Extract text using OCR
        text = ocr_tif_file(reader, file_path)
        
        if text and "[Error" not in text:
            print("OK")
            # Add OCR'd text
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(11)
        else:
            print("NO TEXT")
            doc.add_paragraph("[No text detected in image]")
        
        # Add separator
        doc.add_paragraph('_' * 80)
        
        # Add page break
        if idx < len(tif_files):
            doc.add_page_break()
    
    # Save document
    try:
        doc.save(output_path)
        file_size = os.path.getsize(output_path) / 1024 / 1024
        print(f"\nDocument saved to: {output_path}")
        print(f"File size: {file_size:.2f} MB")
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False

def main():
    """Main execution function"""
    
    tif_folder = r"c:\xampp\htdocs\projectx\watanybot\watany_kb\docs\09_archive_tif_ocr"
    output_path = r"c:\xampp\htdocs\projectx\watanybot\watany_kb\docs\09_archive_tif_ocr\OCR_Output.docx"
    
    print("=" * 80)
    print("TIF File OCR to Word Document Converter")
    print("=" * 80)
    
    if not os.path.isdir(tif_folder):
        print(f"Folder not found: {tif_folder}")
        return False
    
    print("\nCreating Word document from TIF files...\n")
    success = create_word_document(tif_folder, output_path)
    
    print("\n" + "=" * 80)
    if success:
        print("OCR conversion completed successfully!")
    else:
        print("OCR conversion failed")
    print("=" * 80)
    
    return success

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
