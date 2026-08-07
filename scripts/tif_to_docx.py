#!/usr/bin/env python3
"""
Fast TIF to Word Document Converter
Creates a Word document with embedded TIF images and basic metadata
"""

import os
import sys
from pathlib import Path
from datetime import datetime

def create_word_with_images(tif_folder, output_path):
    """Create a Word document with embedded TIF images"""
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Installing python-docx...")
        os.system("pip install python-docx -q")
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Get all TIF files
    tif_files = sorted([f for f in os.listdir(tif_folder) if f.lower().endswith('.tif')])
    
    if not tif_files:
        print("No TIF files found!")
        return False
    
    print(f"Found {len(tif_files)} TIF files")
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Document Scan Archive', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata section
    doc.add_heading('Document Information', level=2)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    rows = table.rows
    rows[0].cells[0].text = "Source"
    rows[0].cells[1].text = tif_folder
    rows[1].cells[0].text = "Total Pages"
    rows[1].cells[1].text = str(len(tif_files))
    rows[2].cells[0].text = "Generated"
    rows[2].cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    rows[3].cells[0].text = "Document Type"
    rows[3].cells[1].text = "Scanned Documents (TIFF Format)"
    rows[4].cells[0].text = "Format"
    rows[4].cells[1].text = "Embedded Images"
    
    doc.add_page_break()
    
    # Process each TIF file
    print("\nAdding images to document...\n")
    for idx, tif_file in enumerate(tif_files, 1):
        file_path = os.path.join(tif_folder, tif_file)
        print(f"[{idx}/{len(tif_files)}] {tif_file}...", end=" ", flush=True)
        
        try:
            # Add page heading
            page_heading = doc.add_heading(f'Page {idx} - {tif_file}', level=2)
            
            # Add file info
            file_stats = os.stat(file_path)
            info = f"File: {tif_file} | Size: {file_stats.st_size / 1024:.1f} KB | Modified: {datetime.fromtimestamp(file_stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S')}"
            doc.add_paragraph(info, style='Caption')
            
            # Embed the image
            try:
                doc.add_picture(file_path, width=Inches(6.0))
                print("OK")
            except Exception as e:
                print(f"ERROR: {e}")
                doc.add_paragraph(f"[Could not embed image: {str(e)}]")
            
            doc.add_paragraph()
            
            # Add page break except on last page
            if idx < len(tif_files):
                doc.add_page_break()
        
        except Exception as e:
            print(f"ERROR: {e}")
            doc.add_paragraph(f"[Error processing {tif_file}: {str(e)}]")
    
    # Save document
    try:
        doc.save(output_path)
        file_size = os.path.getsize(output_path) / 1024 / 1024
        print(f"\nDocument saved!")
        print(f"Output: {output_path}")
        print(f"Size: {file_size:.2f} MB")
        return True
    except Exception as e:
        print(f"Error saving: {e}")
        return False

def main():
    """Main execution"""
    
    tif_folder = r"c:\xampp\htdocs\projectx\watanybot\watany_kb\docs\09_archive_tif_ocr"
    output_path = r"c:\xampp\htdocs\projectx\watanybot\watany_kb\docs\09_archive_tif_ocr\Scanned_Document.docx"
    
    print("=" * 70)
    print("TIF to Word Document Converter")
    print("=" * 70)
    print()
    
    if not os.path.isdir(tif_folder):
        print(f"Error: Folder not found: {tif_folder}")
        return False
    
    success = create_word_with_images(tif_folder, output_path)
    
    print("\n" + "=" * 70)
    if success:
        print("SUCCESS: Document created!")
    else:
        print("ERROR: Document creation failed")
    print("=" * 70)
    
    return success

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
