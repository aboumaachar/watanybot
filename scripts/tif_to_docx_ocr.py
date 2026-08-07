#!/usr/bin/env python3
"""
TIF to Word with OCR Text Extraction
Creates a Word document with TIF images and extracted text using Tesseract OCR
"""

import os
import sys
import subprocess
from datetime import datetime

def sanitize_for_xml(text):
    """Remove control characters incompatible with XML"""
    if not text:
        return text
    # Remove NULL bytes and other control characters
    # Keep only valid XML characters
    valid_chars = []
    for char in text:
        code = ord(char)
        # Keep printable characters and common whitespace
        if (0x20 <= code <= 0xD7FF or 
            code in [0x9, 0xA, 0xD] or  # tab, newline, carriage return
            0xE000 <= code <= 0xFFFD or
            0x10000 <= code <= 0x10FFFF):
            valid_chars.append(char)
    return ''.join(valid_chars)

def check_tesseract():
    """Check if Tesseract OCR is installed"""
    try:
        result = subprocess.run(['tesseract', '--version'], 
                            capture_output=True, text=True)
        if result.returncode == 0:
            return True
    except FileNotFoundError:
        pass
    
    # Check common installation paths
    paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    
    for path in paths:
        if os.path.exists(path):
            return path
    
    return False

def ocr_with_tesseract(image_path, tesseract_cmd=None):
    """Extract text using Tesseract OCR"""
    try:
        if tesseract_cmd:
            cmd = [tesseract_cmd, image_path, "stdout"]
        else:
            cmd = ['tesseract', image_path, "stdout"]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        return result.stdout if result.returncode == 0 else ""
    except Exception as e:
        return ""

def create_word_with_ocr(tif_folder, output_path):
    """Create a Word document with images and OCR text"""
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from PIL import Image
    except ImportError:
        print("Installing python-docx and pillow...")
        os.system("pip install python-docx pillow -q")
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from PIL import Image
    
    # Get all TIF files
    tif_files = sorted([f for f in os.listdir(tif_folder) if f.lower().endswith('.tif')])
    
    if not tif_files:
        print("No TIF files found!")
        return False
    
    print(f"Found {len(tif_files)} TIF files")
    
    # Check for Tesseract
    tesseract_cmd = check_tesseract()
    if tesseract_cmd is True:
        tesseract_cmd = 'tesseract'
        print("Tesseract OCR found on PATH")
    elif tesseract_cmd:
        print(f"Tesseract OCR found at: {tesseract_cmd}")
    else:
        print("Tesseract OCR not found - creating document with images only")
        tesseract_cmd = None
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Scanned Document Archive with OCR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata section
    doc.add_heading('Archive Information', level=2)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    rows = table.rows
    rows[0].cells[0].text = "Source Folder"
    rows[0].cells[1].text = tif_folder
    rows[1].cells[0].text = "Total Pages"
    rows[1].cells[1].text = str(len(tif_files))
    rows[2].cells[0].text = "Generated"
    rows[2].cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    rows[3].cells[0].text = "OCR Method"
    rows[3].cells[1].text = "Tesseract" if tesseract_cmd else "Images Only"
    rows[4].cells[0].text = "Format"
    rows[4].cells[1].text = "Embedded TIFF Images"
    rows[5].cells[0].text = "Language"
    rows[5].cells[1].text = "Arabic + English"
    
    doc.add_page_break()
    
    # Table of contents
    doc.add_heading('Table of Contents', level=2)
    for idx, tif_file in enumerate(tif_files, 1):
        doc.add_paragraph(f"Page {idx}: {tif_file}", style='List Number')
    
    doc.add_page_break()
    
    # Process each TIF file
    print("\nProcessing images and extracting text...\n")
    for idx, tif_file in enumerate(tif_files, 1):
        file_path = os.path.join(tif_folder, tif_file)
        print(f"[{idx}/{len(tif_files)}] {tif_file}...", end=" ", flush=True)
        
        try:
            # Add page heading
            page_heading = doc.add_heading(f'Page {idx} - {tif_file}', level=2)
            
            # Add file info
            file_stats = os.stat(file_path)
            file_size_kb = file_stats.st_size / 1024
            mod_time = datetime.fromtimestamp(file_stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
            info = f"File: {tif_file} ({file_size_kb:.1f} KB) | Modified: {mod_time}"
            doc.add_paragraph(info, style='Caption')
            
            # Try to embed the image
            try:
                doc.add_picture(file_path, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f"[Image preview unavailable: {str(e)}]")
            
            # Extract text if Tesseract is available
            if tesseract_cmd:
                text = ocr_with_tesseract(file_path, tesseract_cmd if tesseract_cmd != 'tesseract' else None)
                if text.strip():
                    # Sanitize text for XML compatibility
                    clean_text = sanitize_for_xml(text)
                    if clean_text.strip():
                        doc.add_heading('Extracted Text', level=3)
                        p = doc.add_paragraph(clean_text)
                        for run in p.runs:
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(64, 64, 64)
            
            doc.add_paragraph()
            
            print("OK")
            
            # Add page break except on last page
            if idx < len(tif_files):
                doc.add_page_break()
        
        except Exception as e:
            print(f"ERROR: {e}")
    
    # Save document
    try:
        doc.save(output_path)
        file_size = os.path.getsize(output_path) / 1024 / 1024
        print(f"\nDocument saved!")
        print(f"Output: {output_path}")
        print(f"Size: {file_size:.2f} MB")
        return True
    except Exception as e:
        print(f"Error saving: {e}")
        return False

def main():
    """Main execution"""
    
    tif_folder = r"c:\xampp\htdocs\projectx\watanybot\watany_kb\docs\09_archive_tif_ocr"
    output_path = r"c:\xampp\htdocs\projectx\watanybot\watany_kb\docs\09_archive_tif_ocr\Scanned_Document_with_OCR.docx"
    
    print("=" * 70)
    print("TIF to Word with OCR Converter")
    print("=" * 70)
    print()
    
    if not os.path.isdir(tif_folder):
        print(f"Error: Folder not found: {tif_folder}")
        return False
    
    success = create_word_with_ocr(tif_folder, output_path)
    
    print("\n" + "=" * 70)
    if success:
        print("SUCCESS: Document created!")
    else:
        print("ERROR: Document creation failed")
    print("=" * 70)
    
    return success

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
